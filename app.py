# Import necessary libraries  
import streamlit as st  
import plotly.graph_objects as go  
import numpy as np  
import pandas as pd  
import io  
import base64  
from datetime import datetime  
from pyopenms import MSExperiment, MzMLFile  
import os  
from PIL import Image  
import os  
  
# Install required packages if not already installed  
try:  
    import docx  
except ImportError:  
    st.info("Installing python-docx...")  
    os.system("pip install python-docx")  
    import docx  
  
# Import reportlab components for PDF generation  
from reportlab.lib.pagesizes import letter  
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage  
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  
from reportlab.lib import colors  
from reportlab.lib.units import inch  
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT  
  
# Try to import kaleido for Plotly image export  
try:  
    import kaleido  
    kaleido_available = True  
except ImportError:  
    kaleido_available = False  
    st.warning("Kaleido package is required for image export. Install it using 'pip install -U kaleido'")  
    st.info("Attempting to install kaleido...")  
    os.system("pip install -U kaleido")  
    st.info("Please restart the app after installation.")  
  
# Set page configuration and title  
st.set_page_config(page_title="mzML Chromatogram Viewer", layout="wide")  
st.title("mzML Chromatogram Viewer")  
st.markdown("Upload an mzML file to view its corresponding chromatogram, mass spectra data, and generate a modern PDF report.")  
  
# Sidebar for logo upload and PDF title entry  
with st.sidebar:  
    st.header("Logo & PDF Settings")  
    logo_file = st.file_uploader("Upload Your Logo", type=["png", "jpg", "jpeg"])  
    if logo_file is not None:  
        # Save the logo to a temporary file  
        with open("temp_logo.png", "wb") as f:  
            f.write(logo_file.getbuffer())  
        st.success("Logo uploaded successfully!")  
        st.image("temp_logo.png", width=200)  
    pdf_title_input = st.text_input("Enter the PDF Title", "My mzML Report")  
  
# -----------------------------------------------  
# Helper functions for processing mzML files  
# -----------------------------------------------  
  
def load_mzml_file(file_bytes):  
    exp = MSExperiment()  
    with open("temp.mzML", "wb") as f:  
        f.write(file_bytes)  
    mzml_file = MzMLFile()  
    mzml_file.load("temp.mzML", exp)  
    return exp  
  
def extract_chromatogram(exp):  
    chromatograms = exp.getChromatograms()  
    if len(chromatograms) == 0:  
        return None, None  
<<<<<<< HEAD
    # For simplicity, use the first chromatogram which is assumed to be the TIC.  
    tic = chromatograms[0]  
    # tic.get_peaks() returns a tuple (times, intensities) as numpy arrays  
    peaks = tic.get_peaks()  
    tic_times = peaks[0].tolist()   # Convert numpy array to list  
    tic_intensities = peaks[1].tolist()  
    return tic_times, tic_intensities  
  
def extract_mass_spectra(exp):  
    spectra = exp.getSpectra()  
    if len(spectra) == 0:  
        return pd.DataFrame()  
      
    # Get the first spectrum for demonstration  
    spectrum = spectra[0]  
    peaks = spectrum.get_peaks()  
    mz_values = peaks[0]  
    intensities = peaks[1]  
      
    # Create a DataFrame with the top peaks by intensity  
    df = pd.DataFrame({  
        "m/z": mz_values,  
        "Intensity": intensities  
    })  
    df = df.sort_values(by="Intensity", ascending=False).reset_index(drop=True)  
=======
    chrom = chromatograms[0]  
    peaks = chrom.get_peaks()  
    if isinstance(peaks, tuple):  
        times, intensities = peaks  
    else:  
        times = [p.getRT() for p in peaks]  
        intensities = [p.getIntensity() for p in peaks]  
    return times, intensities  
  
def extract_mass_spectra(exp):  
    spectra = exp.getSpectra()  
    masses = []  
    intensities = []  
    rts = []  
      
    for spectrum in spectra:  
        rt = spectrum.getRT()  
        peaks = spectrum.get_peaks()  
          
        if isinstance(peaks, tuple):  
            mz_values, intensity_values = peaks  
        else:  
            mz_values = [p.getMZ() for p in peaks]  
            intensity_values = [p.getIntensity() for p in peaks]  
          
        for mz, intensity in zip(mz_values, intensity_values):  
            masses.append(mz)  
            intensities.append(intensity)  
            rts.append(rt)  
      
    if not masses:  
        return pd.DataFrame()  
      
    df = pd.DataFrame({  
        'Retention Time (s)': rts,  
        'm/z': masses,  
        'Intensity': intensities  
    })  
      
    # Sort by intensity (descending)  
    df = df.sort_values('Intensity', ascending=False)  
>>>>>>> parent of 3d1de4d (Update app.py)
    return df  
  
def extract_eic(exp, target_mass, tolerance):  
    spectra = exp.getSpectra()  
<<<<<<< HEAD
    if len(spectra) == 0:  
        return None, None  
      
    rt_values = []  
=======
    times = []  
>>>>>>> parent of 3d1de4d (Update app.py)
    intensities = []  
      
    for spectrum in spectra:  
        rt = spectrum.getRT()  
        peaks = spectrum.get_peaks()  
<<<<<<< HEAD
        mz_values = peaks[0]  
        peak_intensities = peaks[1]  
          
        # Find peaks within the mass tolerance  
        for i, mz in enumerate(mz_values):  
            if abs(mz - target_mass) <= tolerance:  
                rt_values.append(rt)  
                intensities.append(peak_intensities[i])  
                break  
        else:  
            # No peak found within tolerance, add 0 intensity  
            rt_values.append(rt)  
            intensities.append(0)  
      
    return rt_values, intensities  
  
# Function to create a DOCX template for Carbone.io  
def create_docx_template():  
    doc = docx.Document()  
      
    # Add title with placeholder  
    doc.add_heading('{title}', level=1)  
      
    # Add date and filename  
    doc.add_paragraph('Date: {date}')  
    doc.add_paragraph('Filename: {filename}')  
      
    # Add placeholder for TIC image  
    p = doc.add_paragraph()  
    p.add_run().add_text('Total Ion Chromatogram:')  
    doc.add_paragraph('{tic_image}')  
      
    # Add placeholder for mass spectra table  
    doc.add_heading('Mass Spectra Data (Top 10)', level=2)  
    doc.add_paragraph('{mass_data}')  
      
    # Add placeholder for EIC image  
    p = doc.add_paragraph()  
    p.add_run().add_text('Extracted Ion Chromatogram:')  
    doc.add_paragraph('{eic_image}')  
      
    # Add footer  
    section = doc.sections[0]  
    footer = section.footer  
    footer_para = footer.paragraphs[0]  
    footer_para.text = '{footer}'  
      
    # Save the template  
    template_path = "generated_template.docx"  
    doc.save(template_path)  
    return template_path  
  
# Function to generate PDF via Carbone.io API  
def generate_pdf_via_carbone(template_file, data, api_key):  
    # Read the DOCX template as binary  
    with open(template_file, "rb") as f:  
        template_data = f.read()  
      
    # Create a multipart/form-data payload for Carbone.io  
    files = {  
        "template": ("template.docx", template_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")  
    }  
      
    # Convert data to JSON  
    data_json = json.dumps(data)  
      
    # Set up the request  
    headers = {  
        "Authorization": f"Bearer {api_key}",  
        "carbone-version": "4"  
    }  
      
    payload = {  
        "data": data_json,  
        "convertTo": "pdf"  
    }  
      
    # Send the request to Carbone.io  
    try:  
        response = requests.post(  
            "https://api.carbone.io/render",  
            headers=headers,  
            files=files,  
            data=payload  
        )  
          
        if response.status_code == 200:  
            # Get the rendered PDF  
            render_id = response.json().get("data", {}).get("renderId")  
            if render_id:  
                # Get the rendered PDF  
                pdf_response = requests.get(  
                    f"https://api.carbone.io/render/{render_id}",  
                    headers=headers  
                )  
                if pdf_response.status_code == 200:  
                    return pdf_response.content  
          
        st.error(f"Error generating PDF: {response.text}")  
        return None  
    except Exception as e:  
        st.error(f"Error connecting to Carbone.io: {str(e)}")  
        return None  
=======
          
        if isinstance(peaks, tuple):  
            mz_values, intensity_values = peaks  
        else:  
            mz_values = [p.getMZ() for p in peaks]  
            intensity_values = [p.getIntensity() for p in peaks]  
          
        # Find peaks within the mass tolerance  
        matching_intensities = [  
            intensity for mz, intensity in zip(mz_values, intensity_values)  
            if abs(mz - target_mass) <= tolerance  
        ]  
          
        if matching_intensities:  
            times.append(rt)  
            intensities.append(sum(matching_intensities))  
        else:  
            times.append(rt)  
            intensities.append(0)  
      
    return times, intensities  
  
def get_download_link(buffer, filename):  
    b64 = base64.b64encode(buffer.getvalue()).decode()  
    return f'<a href="data:application/pdf;base64,{b64}" download="{filename}_report.pdf" class="download-button">Download PDF Report</a>'  
  
def create_pdf_report(filename, tic_fig, eic_fig, mass_df, target_mass, tolerance, pdf_title):  
    pdf_buffer = io.BytesIO()  
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)  
      
    # Create styles  
    styles = getSampleStyleSheet()  
    custom_title_style = ParagraphStyle(  
        name='CustomTitle',  
        fontName='Helvetica-Bold',  
        fontSize=18,  
        leading=22,  
        alignment=TA_CENTER,  
        spaceAfter=12,  
        textColor=colors.HexColor("#2563EB")  
    )  
      
    heading_style = ParagraphStyle(  
        name='CustomHeading',  
        fontName='Helvetica-Bold',  
        fontSize=14,  
        leading=18,  
        spaceAfter=6,  
        textColor=colors.HexColor("#2563EB")  
    )  
      
    normal_style = ParagraphStyle(  
        name='CustomNormal',  
        fontName='Helvetica',  
        fontSize=12,  
        leading=14,  
        spaceAfter=6  
    )  
      
    right_style = ParagraphStyle(  
        name='CustomRight',  
        fontName='Helvetica',  
        fontSize=10,  
        leading=12,  
        alignment=TA_RIGHT  
    )  
      
    # Elements to add to the PDF  
    elements = []  
      
    # Add logo if available  
    if os.path.exists("temp_logo.png"):  
        try:  
            # Get the original image dimensions  
            img = Image.open("temp_logo.png")  
            img_width, img_height = img.size  
              
            # Calculate the aspect ratio  
            aspect_ratio = img_height / img_width  
              
            # Set a maximum width for the logo (2 inches)  
            max_width = 2 * inch  
              
            # Calculate the height based on the aspect ratio  
            logo_width = min(max_width, img_width)  
            logo_height = logo_width * aspect_ratio  
              
            # Create a table for the header with logo and title  
            logo_img = RLImage("temp_logo.png", width=logo_width, height=logo_height)  
              
            # Create a table for the header  
            header_data = [[logo_img, Paragraph(pdf_title, custom_title_style)]]  
            header_table = Table(header_data, colWidths=[logo_width, 5*inch])  
            header_table.setStyle(TableStyle([  
                ('ALIGN', (0, 0), (0, 0), 'LEFT'),  
                ('ALIGN', (1, 0), (1, 0), 'RIGHT'),  
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  
            ]))  
            elements.append(header_table)  
        except Exception as e:  
            # If there's an error with the logo, just add the title  
            elements.append(Paragraph(pdf_title, custom_title_style))  
            elements.append(Paragraph(f"Logo error: {str(e)}", normal_style))  
    else:  
        # No logo, just add the title  
        elements.append(Paragraph(pdf_title, custom_title_style))  
      
    elements.append(Spacer(1, 20))  
      
    # Add file information  
    elements.append(Paragraph(f"<b>File:</b> {filename}", normal_style))  
    elements.append(Paragraph(f"<b>Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))  
    elements.append(Spacer(1, 20))  
      
    # Add TIC plot  
    elements.append(Paragraph("<b>Total Ion Chromatogram</b>", heading_style))  
    elements.append(Spacer(1, 10))  
      
    # Export TIC plot as image  
    if kaleido_available:  
        try:  
            tic_fig.write_image("temp_tic.png", scale=2)  
            tic_img = RLImage("temp_tic.png", width=6.5*inch, height=4*inch)  
            elements.append(tic_img)  
        except Exception as e:  
            elements.append(Paragraph(f"Error generating TIC image: {str(e)}", normal_style))  
    else:  
        elements.append(Paragraph("Kaleido package not available. Cannot include plot images.", normal_style))  
      
    elements.append(Spacer(1, 20))  
      
    # Add EIC plot  
    elements.append(Paragraph(f"<b>Extracted Ion Chromatogram (m/z {target_mass} ± {tolerance})</b>", heading_style))  
    elements.append(Spacer(1, 10))  
      
    # Export EIC plot as image  
    if kaleido_available:  
        try:  
            eic_fig.write_image("temp_eic.png", scale=2)  
            eic_img = RLImage("temp_eic.png", width=6.5*inch, height=4*inch)  
            elements.append(eic_img)  
        except Exception as e:  
            elements.append(Paragraph(f"Error generating EIC image: {str(e)}", normal_style))  
      
    elements.append(Spacer(1, 20))  
      
    # Add mass spectra data table  
    if not mass_df.empty:  
        data = [mass_df.columns.to_list()] + mass_df.values.tolist()  
        mass_table = Table(data, hAlign='CENTER')  
        mass_table.setStyle(TableStyle([  
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#2563EB")),  
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),  
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),  
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),  
            ('FONTSIZE', (0,0), (-1,0), 12),  
            ('BOTTOMPADDING', (0,0), (-1,0), 8),  
            ('BACKGROUND', (0,1), (-1,-1), colors.whitesmoke),  
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey)  
        ]))  
        elements.append(Paragraph("<b>Top 10 Mass Spectra Peaks</b>", heading_style))  
        elements.append(Spacer(1, 10))  
        elements.append(mass_table)  
        elements.append(Spacer(1, 20))  
      
    # Additional details  
    elements.append(Paragraph(f"Extracted Ion Chromatogram generated for target mass <b>{target_mass}</b> ± <b>{tolerance}</b> m/z.", normal_style))  
    elements.append(Spacer(1, 30))  
      
    # Footer  
    elements.append(Paragraph("© 2025 Kapelczak Metabolomics", right_style))  
      
    doc.build(elements)  
    pdf_buffer.seek(0)  
    return pdf_buffer  
>>>>>>> parent of 3d1de4d (Update app.py)
  
# -----------------------------------------------  
# Main Streamlit app  
# -----------------------------------------------  
uploaded_file = st.file_uploader("Choose an mzML file", type=["mzML"])  
  
if uploaded_file is not None:  
    file_bytes = uploaded_file.read()  
    with st.spinner("Loading mzML file..."):  
        experiment = load_mzml_file(file_bytes)  
    st.success("File loaded successfully!")  
      
    tic_times, tic_intensities = extract_chromatogram(experiment)  
    if tic_times is None or tic_intensities is None:  
        st.error("No chromatogram data found in this mzML file.")  
    else:  
        show_points = st.checkbox("Show individual data points on TIC", value=True, key="tic_toggle")  
        mode = "lines+markers" if show_points else "lines"  
      
        tic_fig = go.Figure()  
        tic_fig.add_trace(go.Scatter(  
            x=tic_times,  
            y=tic_intensities,  
            mode=mode,  
            line=dict(color="#2563EB"),  
            marker=dict(color="#D324EB", size=6)  
        ))  
        tic_fig.update_layout(  
            title=dict(text="Total Ion Chromatogram", x=0.5, xanchor="center", font=dict(size=20, color="#171717")),  
            xaxis_title="Retention Time (s)",  
            yaxis_title="Intensity",  
            plot_bgcolor="#FFFFFF",  
            paper_bgcolor="#FFFFFF"  
        )  
        st.plotly_chart(tic_fig)  
      
        # Extract Mass Spectra  
        mass_df = extract_mass_spectra(experiment)  
        if mass_df.empty:  
            st.warning("No mass spectra data available.")  
        else:  
            st.subheader("Mass Spectra Data (Top 10)")  
            st.dataframe(mass_df.head(10))  
          
        # EIC extraction settings  
<<<<<<< HEAD
        st.markdown("### Extract Ion Chromatogram")  
=======
        st.markdown("### Extract Ion Chromatogram (EIC)")  
>>>>>>> parent of 3d1de4d (Update app.py)
        col1, col2 = st.columns(2)  
        with col1:  
            target_mass = st.number_input("Target m/z", value=400.0, min_value=0.0, format="%.4f")  
        with col2:  
<<<<<<< HEAD
            tolerance = st.number_input("Tolerance (±)", value=0.5, min_value=0.0001, format="%.4f")  
          
        eic_times, eic_intensities = extract_eic(experiment, target_mass, tolerance)  
        if eic_times is None or eic_intensities is None or len(eic_times) == 0:  
            st.warning(f"No data found for m/z {target_mass} ± {tolerance}")  
            eic_fig = None  
        else:  
            show_eic_points = st.checkbox("Show individual data points on EIC", value=True, key="eic_toggle")  
            eic_mode = "lines+markers" if show_eic_points else "lines"  
              
            eic_fig = go.Figure()  
            eic_fig.add_trace(go.Scatter(  
                x=eic_times,  
                y=eic_intensities,  
                mode=eic_mode,  
                line=dict(color="#24EB84"),  
                marker=dict(color="#B2EB24", size=6)  
            ))  
            eic_fig.update_layout(  
                title=dict(text=f"Extracted Ion Chromatogram (m/z {target_mass} ± {tolerance})",   
                          x=0.5, xanchor="center", font=dict(size=20, color="#171717")),  
                xaxis_title="Retention Time (s)",  
                yaxis_title="Intensity",  
                plot_bgcolor="#FFFFFF",  
                paper_bgcolor="#FFFFFF"  
            )  
            st.plotly_chart(eic_fig)  
          
        # Generate PDF Report  
        if st.button("Generate PDF Report"):  
            with st.spinner("Preparing report data..."):  
                # Convert plots to base64 images  
                tic_png_buffer = io.BytesIO()  
                tic_fig.write_image(tic_png_buffer, format="png", scale=2)  
                tic_png_buffer.seek(0)  
                tic_png_b64 = base64.b64encode(tic_png_buffer.read()).decode('utf-8')  
                  
                eic_png_b64 = None  
                if eic_fig is not None:  
                    eic_png_buffer = io.BytesIO()  
                    eic_fig.write_image(eic_png_buffer, format="png", scale=2)  
                    eic_png_buffer.seek(0)  
                    eic_png_b64 = base64.b64encode(eic_png_buffer.read()).decode('utf-8')  
                  
                # Convert logo to base64 if available  
                logo_data_b64 = None  
                if os.path.exists("temp_logo.png"):  
                    with open("temp_logo.png", "rb") as f:  
                        logo_data_b64 = base64.b64encode(f.read()).decode('utf-8')  
                  
                # Create data payload for Carbone.io  
                data_payload = {  
                    "title": pdf_title_input,  
                    "date": datetime.now().strftime("%Y-%m-%d"),  
                    "filename": uploaded_file.name,  
                    "tic_image": tic_png_b64,  
                    "eic_image": eic_png_b64 if eic_png_b64 else "",  
                    "mass_data": mass_df.head(10).to_dict(orient="records"),  
                    "footer": "© 2025 Kapelczak Metabolomics",  
                    "logo": logo_data_b64 if logo_data_b64 else ""  
                }  
                  
                # Create a DOCX template  
                template_file = create_docx_template()  
                  
                with st.spinner("Generating PDF report via Carbone.io..."):  
                    pdf_content = generate_pdf_via_carbone(template_file, data_payload, CARBONE_API_KEY)  
                  
                if pdf_content is not None:  
                    st.success("PDF report generated successfully!")  
                    # Provide a download link for the generated PDF  
                    b64_pdf = base64.b64encode(pdf_content).decode('utf-8')  
                    href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="{uploaded_file.name.split(".")[0]}_report.pdf"><div class="download-button">Download PDF Report</div></a>'  
                    st.markdown("""  
                        <style>  
                        .download-button {  
                            display: inline-block;  
                            padding: 0.75em 1.5em;  
                            color: white;  
                            background-color: #2563EB;  
                            border-radius: 6px;  
                            text-decoration: none;  
                            font-weight: bold;  
                            margin-top: 15px;  
                            box-shadow: 0 4px 6px rgba(37, 99, 235, 0.2);  
                            transition: all 0.3s ease;  
                        }  
                        .download-button:hover {  
                            background-color: #1D4ED8;  
                            box-shadow: 0 6px 8px rgba(37, 99, 235, 0.3);  
                            transform: translateY(-2px);  
                        }  
                        </style>  
                    """, unsafe_allow_html=True)  
                    st.markdown(href, unsafe_allow_html=True)  
=======
            tolerance = st.number_input("Mass Tolerance (±)", value=0.5, min_value=0.0001, format="%.4f")  
          
        # Extract EIC  
        eic_times, eic_intensities = extract_eic(experiment, target_mass, tolerance)  
          
        # Plot EIC  
        eic_fig = go.Figure()  
        eic_fig.add_trace(go.Scatter(  
            x=eic_times,  
            y=eic_intensities,  
            mode="lines",  
            line=dict(color="#24EB84")  
        ))  
        eic_fig.update_layout(  
            title=dict(text=f"Extracted Ion Chromatogram (m/z {target_mass} ± {tolerance})", x=0.5, xanchor="center", font=dict(size=20, color="#171717")),  
            xaxis_title="Retention Time (s)",  
            yaxis_title="Intensity",  
            plot_bgcolor="#FFFFFF",  
            paper_bgcolor="#FFFFFF"  
        )  
        st.plotly_chart(eic_fig)  
          
        # PDF Report Generation  
        st.markdown("### PDF Report")  
        if st.button("Generate PDF Report"):  
            with st.spinner("Generating PDF report..."):  
                # Create PDF report  
                pdf_buffer = create_pdf_report(  
                    filename=uploaded_file.name,  
                    tic_fig=tic_fig,  
                    eic_fig=eic_fig,  
                    mass_df=mass_df.head(10),  
                    target_mass=target_mass,  
                    tolerance=tolerance,  
                    pdf_title=pdf_title_input  
                )  
                  
                # Create download link  
                download_link = get_download_link(pdf_buffer, uploaded_file.name.split('.')[0])  
                  
                # Display download link with custom styling  
                st.markdown("""  
                <style>  
                .download-button {  
                    display: inline-block;  
                    padding: 0.75em 1.5em;  
                    color: white;  
                    background-color: #2563EB;  
                    border-radius: 6px;  
                    text-decoration: none;  
                    font-weight: bold;  
                    margin-top: 15px;  
                    box-shadow: 0 4px 6px rgba(37, 99, 235, 0.2);  
                    transition: all 0.3s ease;  
                }  
                .download-button:hover {  
                    background-color: #1D4ED8;  
                    box-shadow: 0 6px 8px rgba(37, 99, 235, 0.3);  
                    transform: translateY(-2px);  
                }  
                </style>  
                """, unsafe_allow_html=True)  
                  
                st.markdown(download_link, unsafe_allow_html=True)  
                st.success("PDF report generated successfully! Click the button above to download.")  
>>>>>>> parent of 3d1de4d (Update app.py)
else:  
    st.info("Please upload an mzML file to begin.")  
